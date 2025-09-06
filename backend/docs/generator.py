"""Utility functions for generating documents using python-docx and reportlab.

The functions here are intentionally lightweight but demonstrate how dynamic
content can be injected into document templates.
"""
from pathlib import Path
from typing import Dict


# These imports are optional during runtime; the tests will skip if the
# dependencies are not available in the execution environment.
try:  # pragma: no cover - import guarding
    from docx import Document
except Exception:  # pragma: no cover - handled in tests
    Document = None

try:  # pragma: no cover
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:  # pragma: no cover
    canvas = None
    A4 = None


def _ensure_dependencies():
    if Document is None or canvas is None:
        raise RuntimeError("Required libraries 'python-docx' and 'reportlab' are not installed")


def generate_certificate_docx(data: Dict[str, str], output_path: Path) -> Path:
    """Generate a certificate DOCX file using :mod:`python-docx`.

    Parameters
    ----------
    data: dict
        Mapping containing ``name``, ``course`` and ``date``.
    output_path: :class:`pathlib.Path`
        Location where the document will be written.
    """
    _ensure_dependencies()
    doc = Document()
    doc.add_heading("Удостоверение", 0)
    doc.add_paragraph(f"ФИО: {data['name']}")
    doc.add_paragraph(f"Курс: {data['course']}")
    doc.add_paragraph(f"Дата: {data['date']}")
    doc.save(output_path)
    return output_path


def generate_protocol_docx(data: Dict[str, str], output_path: Path) -> Path:
    _ensure_dependencies()
    doc = Document()
    doc.add_heading("Протокол", 0)
    doc.add_paragraph(f"Номер: {data['number']}")
    doc.add_paragraph(f"Дата: {data['date']}")
    doc.add_paragraph(data.get('body', ''))
    doc.save(output_path)
    return output_path


def generate_act_docx(data: Dict[str, str], output_path: Path) -> Path:
    _ensure_dependencies()
    doc = Document()
    doc.add_heading("Акт", 0)
    doc.add_paragraph(f"Номер: {data['number']}")
    doc.add_paragraph(f"Дата: {data['date']}")
    doc.add_paragraph(data.get('body', ''))
    doc.save(output_path)
    return output_path


def generate_certificate_pdf(data: Dict[str, str], output_path: Path) -> Path:
    """Generate a PDF certificate using :mod:`reportlab`."""
    _ensure_dependencies()
    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    c.setFont("Helvetica", 24)
    c.drawCentredString(width / 2, height - 100, "Удостоверение")
    c.setFont("Helvetica", 14)
    c.drawString(100, height - 150, f"ФИО: {data['name']}")
    c.drawString(100, height - 170, f"Курс: {data['course']}")
    c.drawString(100, height - 190, f"Дата: {data['date']}")
    c.showPage()
    c.save()
    return output_path
